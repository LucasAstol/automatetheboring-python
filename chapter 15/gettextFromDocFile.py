#! python3
# gettextFromDocFile.py - gets the text without styles form a word document.

import docx, os
from pathlib import Path


def printTextFromDoc(filePath):
    theDoc = docx.Document(Path(filePath))
    for i in range(len(theDoc.paragraphs)):
        print(theDoc.paragraphs[i].text)